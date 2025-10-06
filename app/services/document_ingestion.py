"""Document ingestion service for handling file uploads and initial processing."""

import os
import uuid
import tempfile
from pathlib import Path
from typing import BinaryIO, Tuple, Optional

import fitz  # PyMuPDF
from fastapi import UploadFile, HTTPException

from app.models.document import (
    Document, 
    DocumentMetadata, 
    DocumentFormat, 
    ProcessingStatus,
    DocumentContent
)
from app.services.enhanced_document_processor import EnhancedDocumentProcessor


class DocumentIngestionService:
    """Service for ingesting and processing uploaded documents."""
    
    SUPPORTED_FORMATS = {
        ".pdf": DocumentFormat.PDF,
        ".txt": DocumentFormat.TEXT,
        ".docx": DocumentFormat.DOCX
    }
    
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB limit
    
    def __init__(self, upload_directory: str = "data/uploads", enhanced_processor: Optional[EnhancedDocumentProcessor] = None, storage_service=None):
        """Initialize the ingestion service.
        
        Args:
            upload_directory: Directory to store uploaded files
            enhanced_processor: Enhanced document processor for LLM integration
            storage_service: Storage service to use (if None, creates new one)
        """
        self.upload_directory = Path(upload_directory)
        self.upload_directory.mkdir(parents=True, exist_ok=True)
        self.enhanced_processor = enhanced_processor
        self.storage_service = storage_service  # Use the provided storage service
    
    async def ingest_document(self, file: UploadFile) -> Document:
        """Ingest an uploaded document and create a Document object.
        
        Args:
            file: The uploaded file
            
        Returns:
            Document object with metadata and initial processing status
            
        Raises:
            HTTPException: If file validation fails
        """
        # Validate file
        await self._validate_file(file)
        
        # Generate unique document ID
        document_id = str(uuid.uuid4())
        
        # Create metadata
        file_extension = Path(file.filename).suffix.lower()
        format_type = self.SUPPORTED_FORMATS[file_extension]
        
        metadata = DocumentMetadata(
            filename=file.filename,
            file_size=file.size,
            format=format_type
        )
        
        # Save file to disk
        file_path = await self._save_file(file, document_id)
        
        # Create document object
        document = Document(
            id=document_id,
            metadata=metadata,
            status=ProcessingStatus.PENDING
        )
        
        # CRITICAL FIX: Always store document first (before processing)
        # Use the storage service passed to us, don't create a new one
        if not hasattr(self, 'storage_service') or self.storage_service is None:
            from app.services.document_storage import DocumentStorageService
            self.storage_service = DocumentStorageService()
        
        # Store the document initially
        await self.storage_service.store_document(document)
        print(f"📦 Document {document.id[:8]} stored initially")
        
        # Process the document content
        try:
            document.status = ProcessingStatus.PROCESSING
            await self.storage_service.update_document(document.id, document)
            
            content = await self._extract_content(file_path, format_type)
            document.content = content
            
            # Use enhanced processor if available
            if self.enhanced_processor:
                document = await self.enhanced_processor.process_document(document)
            else:
                document.status = ProcessingStatus.COMPLETED
                await self.storage_service.update_document(document.id, document)
                
        except Exception as e:
            print(f"❌ Processing failed for {document.id[:8]}: {str(e)}")
            document.status = ProcessingStatus.FAILED
            document.error_message = str(e)
            # CRITICAL: Still store failed documents
            await self.storage_service.update_document(document.id, document)
        
        return document
    
    async def _validate_file(self, file: UploadFile) -> None:
        """Validate uploaded file.
        
        Args:
            file: The uploaded file
            
        Raises:
            HTTPException: If validation fails
        """
        if not file.filename:
            raise HTTPException(status_code=400, detail="No filename provided")
        
        if file.size > self.MAX_FILE_SIZE:
            raise HTTPException(
                status_code=413, 
                detail=f"File too large. Maximum size: {self.MAX_FILE_SIZE / 1024 / 1024}MB"
            )
        
        file_extension = Path(file.filename).suffix.lower()
        if file_extension not in self.SUPPORTED_FORMATS:
            supported = ", ".join(self.SUPPORTED_FORMATS.keys())
            raise HTTPException(
                status_code=415,
                detail=f"Unsupported file format. Supported formats: {supported}"
            )
    
    async def _save_file(self, file: UploadFile, document_id: str) -> Path:
        """Save uploaded file to disk.
        
        Args:
            file: The uploaded file
            document_id: Unique document identifier
            
        Returns:
            Path to saved file
        """
        file_extension = Path(file.filename).suffix.lower()
        file_path = self.upload_directory / f"{document_id}{file_extension}"
        
        # Reset file pointer
        await file.seek(0)
        
        # Save file
        with open(file_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        return file_path
    
    async def _extract_content(self, file_path: Path, format_type: DocumentFormat) -> DocumentContent:
        """Extract text content from the document.
        
        Args:
            file_path: Path to the document file
            format_type: Format of the document
            
        Returns:
            DocumentContent with extracted text
            
        Raises:
            Exception: If content extraction fails
        """
        if format_type == DocumentFormat.PDF:
            return await self._extract_pdf_content(file_path)
        elif format_type == DocumentFormat.TEXT:
            return await self._extract_text_content(file_path)
        elif format_type == DocumentFormat.DOCX:
            return await self._extract_docx_content(file_path)
        else:
            raise NotImplementedError(f"Content extraction for {format_type} not implemented yet")
    
    async def _extract_pdf_content(self, file_path: Path) -> DocumentContent:
        """Extract content from PDF file using PyMuPDF with fallback methods.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            DocumentContent with extracted text
        """
        try:
            doc = fitz.open(file_path)
            raw_text = ""
            page_count = len(doc)  # Store page count BEFORE closing
            extraction_method = "pymupdf_default"
            
            # Try default extraction first
            for page_num in range(page_count):
                page = doc.load_page(page_num)
                page_text = page.get_text()
                raw_text += page_text
                raw_text += "\n\n"  # Page separator
            
            # Check if extraction failed (mostly empty or only special characters)
            printable_chars = sum(1 for c in raw_text if c.isprintable() and not c.isspace())
            if printable_chars < 50 or raw_text.count('`') > printable_chars * 0.5:
                print("⚠️  Default PDF extraction failed, trying alternative methods...")
                raw_text = ""
                
                # Try HTML extraction (better for complex fonts)
                for page_num in range(page_count):
                    page = doc.load_page(page_num)
                    try:
                        html_text = page.get_text("html")
                        # Extract text from HTML (simple method)
                        import re
                        clean_text = re.sub(r'<[^>]+>', ' ', html_text)
                        clean_text = re.sub(r'\s+', ' ', clean_text)
                        if len(clean_text.strip()) > len(raw_text):
                            raw_text = clean_text
                            extraction_method = "pymupdf_html"
                            break
                    except:
                        continue
                
                # If still poor extraction, try JSON method
                if printable_chars < 50:
                    for page_num in range(page_count):
                        page = doc.load_page(page_num)
                        try:
                            import json
                            json_data = json.loads(page.get_text("json"))
                            text_parts = []
                            
                            for block in json_data.get("blocks", []):
                                for line in block.get("lines", []):
                                    for span in line.get("spans", []):
                                        text = span.get("text", "")
                                        if text and text != "`":
                                            text_parts.append(text)
                            
                            if text_parts:
                                page_text = " ".join(text_parts)
                                raw_text += page_text + "\n\n"
                                extraction_method = "pymupdf_json"
                        except:
                            continue
            
            doc.close()  # Close document after extraction
            
            # Basic text cleaning
            cleaned_text = self._clean_text(raw_text)
            
            # Add warning if extraction is still poor
            final_printable = sum(1 for c in cleaned_text if c.isprintable() and not c.isspace())
            if final_printable < 100:
                print(f"⚠️  PDF extraction yielded minimal text ({final_printable} chars). Document may be scanned or have font encoding issues.")
            
            return DocumentContent(
                raw_text=raw_text,
                cleaned_text=cleaned_text,
                page_count=page_count,
                processing_metadata={
                    "extraction_method": extraction_method,
                    "printable_chars": final_printable,
                    "extraction_quality": "good" if final_printable > 500 else "poor" if final_printable > 50 else "failed"
                }
            )
            
        except Exception as e:
            raise Exception(f"Failed to extract PDF content: {str(e)}")
    
    async def _extract_text_content(self, file_path: Path) -> DocumentContent:
        """Extract content from text file.
        
        Args:
            file_path: Path to text file
            
        Returns:
            DocumentContent with text
        """
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                raw_text = f.read()
            
            cleaned_text = self._clean_text(raw_text)
            
            return DocumentContent(
                raw_text=raw_text,
                cleaned_text=cleaned_text,
                processing_metadata={"extraction_method": "direct_read"}
            )
            
        except Exception as e:
            raise Exception(f"Failed to extract text content: {str(e)}")
    
    async def _extract_docx_content(self, file_path: Path) -> DocumentContent:
        """Extract content from DOCX file using python-docx.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            DocumentContent with extracted text
        """
        try:
            from docx import Document
            
            doc = Document(file_path)
            raw_text = ""
            
            # Extract text from all paragraphs
            for paragraph in doc.paragraphs:
                raw_text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        raw_text += cell.text + " "
                    raw_text += "\n"
            
            # Basic text cleaning
            cleaned_text = self._clean_text(raw_text)
            
            return DocumentContent(
                raw_text=raw_text,
                cleaned_text=cleaned_text,
                page_count=1,  # DOCX doesn't have pages in the same way as PDF
                processing_metadata={"extraction_method": "python-docx"}
            )
            
        except ImportError:
            raise Exception("python-docx library is required for DOCX files. Install with: pip install python-docx")
        except Exception as e:
            raise Exception(f"Failed to extract DOCX content: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Strip whitespace and remove empty lines
            line = line.strip()
            if line:
                cleaned_lines.append(line)
        
        # Join with single newlines
        cleaned_text = '\n'.join(cleaned_lines)
        
        # Remove excessive spaces
        import re
        cleaned_text = re.sub(r' +', ' ', cleaned_text)
        
        return cleaned_text
